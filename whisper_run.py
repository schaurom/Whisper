import whisper
import csv
from docx import Document
from docx.shared import Pt
import time


# Startzeit für Programm-laufzeit
startzeit = time.time()

model = whisper.load_model("tiny", download_root='.\models')

options = {"language": "de", "verbose": "true", "word_timestamps": "true", "append_punctuations": "."}
result = model.transcribe("./audio/audio.mp3", **options)
#result = model.transcribe(r"C:\Users\schau\PycharmProjects\Whisper\whisper\audio.mp3", **options)

#Inhalt in Variable schreiben
result_text = result["text"]
print(result_text)

#----------------------------------------------------------------------------------------------------------------
# CSV
# Definieren Sie den Dateinamen, unter dem Sie die CSV-Datei speichern möchten
dateiname = "./audio/audio_transcript.csv"

# Öffne die Datei im Schreibmodus
# mode='w' --> Datei wird überschrieben
# mode='a' --> append --> neue Daten werden in der bestehenden Datei unten drangehängt
with open(dateiname, mode='w', newline='') as csv_datei:
    csv_schreiber = csv.writer(csv_datei)
    # Text in CSV Datei schreiben
    csv_schreiber.writerow([result_text])
#----------------------------------------------------------------------------------------------------------------
# Word
dateiname = "./audio/audio_transcript.docx"

# erstelle ein neues Word Dokument
dokument = Document()

# Fügen Sie den Text in das Dokument ein und spezifizieren Sie die Schriftgröße
absatz = dokument.add_paragraph()
lauf = absatz.add_run(result_text)
lauf.font.size = Pt(14)  # Ändern Sie die Schriftgröße auf 14 Pt

# Fügen Sie den Text in das Dokument ein
#dokument.add_paragraph(result_text)
# Speichern Sie das Dokument in einer Datei
dokument.save(dateiname)


# Endzeit für Programm-laufzeit
endzeit = time.time()
# Die Laufzeit berechnen (in Sekunden)
laufzeit = endzeit - startzeit
print(f"Die Laufzeit des Programms beträgt {laufzeit:.2f} Sekunden.")