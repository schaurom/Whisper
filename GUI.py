# grafische Oberfläche zum Auswählen von Dateien

#-----------------------------------------------------------------------
# Absprung in whisper.py funktioniert noch nicht
#-----------------------------------------------------------------------

import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import subprocess
from docx import Document


def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("MP3 files", "*.mp3")])
    if file_path:
        selected_file_label.config(text=f"Selected File: {file_path}")


def save_docx():
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
    if file_path:
        try:
            # Create a new Word document
            doc = Document()

            # Add content to the document (you can customize this part)
            doc.add_heading('My Document', 0)
            doc.add_paragraph('This is a sample Word document.')

            # Save the document
            doc.save(file_path)

            messagebox.showinfo("Info", f"Document saved as: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")


def run_whisper():
    try:
        subprocess.run(["python", "whisper.py"])
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")


# Create the main window
root = tk.Tk()
root.title("MP3 File Selector, Word Document Saver, and Whisper Runner")

# Set the width and height of the main window
window_width = 400  # Set your desired width
window_height = 200  # Set your desired height
root.geometry(f"{window_width}x{window_height}")

# Create and configure widgets
browse_button = tk.Button(root, text="Browse MP3 File", command=browse_file)
selected_file_label = tk.Label(root, text="Selected File: None")
save_docx_button = tk.Button(root, text="Save Word Doc", command=save_docx)
run_whisper_button = tk.Button(root, text="Run Whisper", command=run_whisper)

# Pack widgets
browse_button.pack(pady=10)
selected_file_label.pack()
save_docx_button.pack(pady=10)
run_whisper_button.pack(pady=10)

# Start the main loop
root.mainloop()
